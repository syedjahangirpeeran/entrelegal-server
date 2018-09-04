# users/views.py
from rest_framework import generics

from . import models
from . import serializers
from .models import Asset
from django.http import HttpResponse
from rest_framework.response import Response
from rest_framework import status

from django.shortcuts import render
from rest_framework.decorators import api_view


from docx import Document
from docx.shared import Inches
from docx.shared import Pt


class UserListView(generics.ListCreateAPIView):
    queryset = models.CustomUser.objects.all()
    serializer_class = serializers.UserSerializer

@api_view(['GET'])
def get_all_user_asset(request):
    try:
    	assests = Asset.objects.get(user=request.user)
    except:
        return Response(status=status.HTTP_404_NOT_FOUND)
    serializer = AssetSerializer(assests)
    return Response(serializer.data)

@api_view(['POST'])
def add_asset(request):
    asset_type = request.POST.get('asset_type', None)
    description = request.POST.get('description', None)
    recipient = request.POST.get('recipient', None)
    relation = request.POST.get('relation', None)
    if not asset_type or not description or not recipient or not relation:
    	error_message = "Missing parameters in request. Send asset_type, description, recipient, relation"
    	return Response(error_message, status=status.HTTP_400_BAD_REQUEST)
    Asset.objects.create(asset_type=asset_type, description=description, recipient=recipient, relation=relation,user=request.user)
    success_message = "Sucessfully added new assest."
    return Response(success_message, status=status.HTTP_201_CREATED)

@api_view(['POST','GET'])
def download(request):
	if request.method == 'GET':
		print(request.GET)
	#get Credentials
		name = str(request.GET.get("name"))
		fatherName = str(request.GET.get("fatherName"))
		address = str(request.GET.get("address"))
		cityName = str(request.GET.get("cityName"))
		pinCode = str(request.GET.get("pinCode"))
		dOB = str(request.GET.get("dOB"))
		sex = str(request.GET.get("sex"))
		martialStatus = str(request.GET.get("martialStatus"))
		religion = str(request.GET.get("religion"))
		spouseName = str(request.GET.get("spouseName"))
		childrenNames = str(request.GET.get("childrenNames"))
		recipientName = str(request.GET.get("recipientName"))
		recipientFatherName = str(request.GET.get("recipientFatherName"))
		recipientRelation = str(request.GET.get("recipientRelation"))
		aRecipientName = str(request.GET.get("aRecipientName"))
		aRecipientFatherName = str(request.GET.get("aRecipientFatherName"))
		aRecipientRelation = str(request.GET.get("aRecipientRelation"))
		executorName = str(request.GET.get("executorName"))
		executorFatherName = str(request.GET.get("executorFatherName"))
		executorRelation = str(request.GET.get("executorRelation"))
 	#Make Document	
		document = Document()
		document.add_heading('Entrelegal Will', 0)
		p = document.add_paragraph("")
		p.add_run('Congratulations! Your Will is now ready to be executed.').bold = True
		
		#document.add_heading('Heading, level 1', level=1)
		#document.add_paragraph('Intense quote', style='IntenseQuote')
		document.add_paragraph('Please read the instructions on this page and the contents of your Will carefully. If you feel that there are certain ambiguities or provisions in your Will that you would want further clarity on, it would be best to have your Will reviewed by a lawyer.');
		x = document.add_paragraph("")
		x.add_run('To make this Will binding, all you are required to do is (a) sign it in the presence of two witnesses of your choice and (b) have the witnesses sign and attest the Will where required. Under Hindu law, the witnesses may also be the beneficiaries under the Will if you so desire.').bold = True
		document.add_paragraph('The details of the witnesses should be filled in where indicated. Ideally, the name, age, fathers name and address of any person named in the Will should be stated clearly. At the very minimum, do indicate the father\'s name. Have the witnesses sign where required, mention the date on which they signed the Will and print their name in block letters.');
		document.add_paragraph('Yet, even then it is advisable to have your Will notaries and/or registered to avoid future disputes, even though it is technically not required under law. A notarised/registered Will can be easier to prove in your absence which is when it would be most required.');
		x = document.add_paragraph('\n')
		x.add_run('THIS IS AN INSTRUCTIONAL PAGE ONLY. YOUR WILL STARTS FROM NEXT PAGE. PLEASE DISCARD THIS PAGE ONCE YOU SIGN YOUR WILL AND GET 2 WITNESSES TO SIGN IT.').bold = True
		document.add_page_break()
	
	#PAGE 2
		document.add_heading('LAST WILL & TESTAMENT OF '+name, level=1)
		x = document.add_paragraph("")
		document.add_paragraph('IN WITNESS WHEREOF I have signed this Will hereunder on the ______ day of _________________, 2018.');
		
		document.add_paragraph('I, '+name+', daughter of Sh. '+fatherName+', resident of '+address+','+cityName+' - '+pinCode+' being of sound mind and memory, do hereby make, publish, and declare this to be my last Will and testament (hereinafter referred to as "this Will" or "my Will") and hereby revoke and render void any former will, testament and codicils thereto (if any).');
		document.add_heading('EXECUTOR NOMINATION', level=1)
		document.add_paragraph('By this Will, I hereby nominate my'+ executorRelation+', '+executorName+', (Father\'s Name:'+executorFatherName +'to be the sole executor of this Will.')
		document.add_heading('MY FAMILY', level=1)
		if martialStatus == "Married":
			document.add_paragraph('At the time of executing this Will, I am married to '+spouseName)
		else:
			document.add_paragraph('At the time of executing this Will, I am Unmarried and single.')
		
		document.add_heading('PRINCIPAL DISTRIBUTION OF MY ASSETS', level=1)
		document.add_paragraph('I wish to leave the entirety of my assets, both movable and immovable, including real estate, mutual funds, stocks, private equity funds, securities, fixed deposits, insurance policies, bank accounts, contents of bank lockers, jewellery, moveable assets, provident funds and any other type of property, assets, investments etc., to my '+recipientRelation+', '+recipientName+' (Father\'s Name: '+ recipientFatherName+').')
		if len(aRecipientName) < 5:
			document.add_paragraph('If such beneficiary predeceases me, then respective share of the predeceased beneficiary shall be distributed to my '+aRecipientRelation+', '+aRecipientName+' (Father\'s Name: '+ aRecipientFatherName+').')
		document.add_heading('DECLARATION', level=1)
		document.add_paragraph('I DECLARE THAT I am executing this Will in complete possession of my mental faculties and in a composed state of mind, out of my own accord without any coercion, compulsion or misrepresentation. I am executing my Will so as to avoid bad blood, disputes, differences and litigation amongst my family members after my death and it is my desire that after my death my Will is acted upon by my family members in its true spirit.')
		document.add_page_break()
# PAGE 3
		document.add_paragraph(name)
		document.add_paragraph('______________________________________')
		document.add_paragraph('NAME: ________________________\n')
		document.add_paragraph('(In Block Letter)\n')
		document.add_heading('AFFIRMATION BY WITNESSES '+name, level=1)
		document.add_paragraph('I, ______________________, (Fathers Name: ________________________), have witnessed the signature of this Will by the Testator, '+name+' and I affirm that the Testator appeared to me to be of sound mind, was not under duress, and the Testator affirmed to me that she was aware of the nature of this Will and signed it freely and voluntarily.')
		document.add_heading('WITNESS I\n', level=1)
		document.add_paragraph('______________________________________\n')
		document.add_paragraph('NAME: ________________________\n')
		document.add_paragraph('(In Block Letter)\n')
		document.add_paragraph('I, ______________________, (Fathers Name: ________________________), have witnessed the signature of this Will by the Testator, '+name+' and I affirm that the Testator appeared to me to be of sound mind, was not under duress, and the Testator affirmed to me that she was aware of the nature of this Will and signed it freely and voluntarily.')
		document.add_heading('WITNESS II\n', level=1)
		document.add_paragraph('______________________________________\n')
		document.add_paragraph('NAME: ________________________\n')
		document.add_paragraph('(In Block Letter)\n')

		response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
		response['Content-Disposition'] = 'attachment; filename=download.docx'
		document.save(response)

		
		return response 

'''	dOB = request.data.get('dOB', None)
	sex = request.data.get('sex', None)
	fatherName = request.data.get('fatherName', None)
	addressN = request.data.get('addressN', None)
	cityName = request.data.get('cityName', None)
	pinCode = request.data.get('pinCode', None)
	martialStatus = request.data.get('martialStatus', None)
	religion = request.data.get('religion', None)
	spouseName = request.data.get('spouseName', None)
	childrenNames = request.data.get('childrenNames', None)
	recipientName = request.data.get('recipientName', None)
	recipientFatherName  = request.data.get('recipientFatherName', None)
	recipientRelation = request.data.get('recipientRelation', None)
	aRecipientName = request.data.get('aRecipientName', None)
	aRecipientFatherName = request.data.get('aRecipientFatherName', None)
	aRecipientRelation = request.data.get('aRecipientRelation', None)
	executorName = request.data.get('executorName', None)
	executorFatherName = request.data.get('executorFatherName', None)
	executorRelation = request.data.get('executorRelation', None)'''